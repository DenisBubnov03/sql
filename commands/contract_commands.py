"""
Модуль для автоматизации формирования договоров для студентов.
"""
import os
from datetime import datetime, date
from docx import Document
from num2words import num2words
from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import ContextTypes
from data_base.db import session
from data_base.models import Student
from commands.states import (
    CONTRACT_MENU, CONTRACT_STUDENT_TG, CONTRACT_TYPE, CONTRACT_ADVANCE_AMOUNT,
    CONTRACT_PAYMENT_TYPE, CONTRACT_MONTHS, CONTRACT_COMMISSION_TYPE,
    CONTRACT_COMMISSION_CUSTOM, CONTRACT_FIO, CONTRACT_ADDRESS, CONTRACT_INN,
    CONTRACT_RS, CONTRACT_KS, CONTRACT_BANK, CONTRACT_BIK, CONTRACT_EMAIL
)
from commands.authorized_users import AUTHORIZED_USERS, NOT_ADMINS
from utils.security import restrict_to


def get_project_root():
    """
    Определяет корневую директорию проекта.
    Ищет директорию, содержащую bot.py или .git, начиная с текущей директории и поднимаясь вверх.
    Если не находит, использует текущую рабочую директорию или директорию на уровень выше commands.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    # Сначала пробуем найти по bot.py или .git
    current_dir = os.path.abspath(os.path.dirname(__file__))
    
    # Поднимаемся вверх по директориям, пока не найдем bot.py или .git
    while current_dir != os.path.dirname(current_dir):  # Пока не достигли корня файловой системы
        if os.path.exists(os.path.join(current_dir, 'bot.py')) or \
           os.path.exists(os.path.join(current_dir, '.git')):
            logger.info(f"Корень проекта найден: {current_dir}")
            return current_dir
        current_dir = os.path.dirname(current_dir)
    
    # Если не нашли, пробуем использовать текущую рабочую директорию
    cwd = os.getcwd()
    if os.path.exists(os.path.join(cwd, 'bot.py')) or \
       os.path.exists(os.path.join(cwd, '.git')):
        logger.info(f"Корень проекта из рабочей директории: {cwd}")
        return cwd
    
    # Последний fallback - директория на уровень выше commands
    fallback_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    logger.warning(f"Корень проекта не найден, используем fallback: {fallback_dir}")
    return fallback_dir

# Шаблоны договоров
CONTRACT_TEMPLATES = {
    "Ручное": "doc/шаблоны/Шаблон ДВОУ Ручное.docx",
    "Авто": "doc/шаблоны/Шаблон ДВОУ Авто .docx",
    "Фуллстек": "doc/шаблоны/Шаблон ДВОУ Фуллстак.docx",
    "Экспресс Авто": "doc/шаблоны/Шаблон Экспресс Авто.docx"
}

# Шаблонные суммы авансового платежа
ADVANCE_AMOUNTS = {
    "Ручное": [
        {"amount": 46000, "text": "Сорок шесть тысяч"},
        {"amount": 45000, "text": "Сорок пять тысяч"},
        {"amount": 41000, "text": "Сорок одна тысяча"}
    ],
    "Авто": [
        {"amount": 86000, "text": "Восемьдесят шесть тысяч"},
        {"amount": 85000, "text": "Восемьдесят пять тысяч"},
        {"amount": 81000, "text": "Восемьдесят одна тысяча"}
    ],
    "Фуллстек": [
        {"amount": 109000, "text": "Сто девять тысяч"},
        {"amount": 108000, "text": "Сто восемь тысяч"},
        {"amount": 104000, "text": "Сто четыре тысячи"}
    ],
    "Экспресс Авто": [
        {"amount": 15999, "text": "Пятнадцать тысяч девятьсот девяносто девять"},
        {"amount": 10999, "text": "Десять тысяч девятьсот девяносто девять"}
    ]
}


def number_to_words_rubles(amount):
    """
    Конвертирует число в пропись на русском языке с добавлением слова "рублей".
    """
    try:
        # Используем num2words для конвертации
        words = num2words(int(amount), lang='ru')
        # Первая буква заглавная
        words = words.capitalize()
        # Добавляем "рублей"
        return f"{words} рублей"
    except Exception as e:
        return f"{amount} рублей"


def get_contract_date_formatted():
    """
    Возвращает дату договора в формате "31 окт. 2025"
    """
    today = date.today()
    months = {
        1: "янв.", 2: "фев.", 3: "мар.", 4: "апр.",
        5: "май", 6: "июн.", 7: "июл.", 8: "авг.",
        9: "сен.", 10: "окт.", 11: "нояб.", 12: "дек."
    }
    return f"{today.day} {months[today.month]} {today.year}"


def get_contract_number():
    """
    Возвращает номер договора в формате "д31/10" где 31/10 - дата создания
    """
    today = date.today()
    return f"{today.day}/{today.month}"

@restrict_to(['admin', 'mentor']) # Разрешаем доступ обеим ролям
async def start_contract_formation(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Начало процесса формирования договора.
    Показывает меню: создать новый или отправить существующий.
    """

    await update.message.reply_text(
        "📄 Формирование договора\n\n"
        "Выберите действие:",
        reply_markup=ReplyKeyboardMarkup(
            [["Создать новый договор"], ["Отправить существующий"], ["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_MENU


async def handle_contract_menu(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает выбор в меню договоров.
    """
    choice = update.message.text.strip()

    if choice == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    if choice == "Создать новый договор":
        await update.message.reply_text(
            "📄 Формирование договора\n\n"
            "Введите Telegram ученика:",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_STUDENT_TG
    elif choice == "Отправить существующий":
        await update.message.reply_text(
            "Введите Telegram ученика для поиска договора:",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        context.user_data['resending_contract'] = True
        return CONTRACT_STUDENT_TG
    else:
        await update.message.reply_text(
            "Выберите действие:",
            reply_markup=ReplyKeyboardMarkup(
                [["Создать новый договор"], ["Отправить существующий"], ["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_MENU


async def handle_student_telegram(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод Telegram ученика и переходит к выбору типа договора.
    """
    telegram = update.message.text.strip()

    if telegram == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    # Нормализуем telegram (убираем @ для поиска, но сохраняем оригинал)
    telegram_clean = telegram.replace("@", "").strip()
    telegram_with_at = f"@{telegram_clean}" if not telegram_clean.startswith("@") else telegram_clean

    # Если отправляем существующий договор
    if context.user_data.get('resending_contract'):
        # Пробуем найти файл с @ и без @
        # Получаем абсолютный путь относительно корня проекта
        base_dir = get_project_root()
        doc_dir = os.path.join(base_dir, "doc")
        
        file_path_with_at = os.path.join(doc_dir, f"{telegram_with_at}.docx")
        file_path_clean = os.path.join(doc_dir, f"{telegram_clean}.docx")

        file_path = None
        if os.path.exists(file_path_with_at):
            file_path = file_path_with_at
        elif os.path.exists(file_path_clean):
            file_path = file_path_clean

        if file_path:
            with open(file_path, 'rb') as doc_file:
                await update.message.reply_document(
                    document=doc_file,
                    filename=os.path.basename(file_path),
                    caption="✅ Договор найден и отправлен!"
                )
            await update.message.reply_text(
                "Договор отправлен.",
                reply_markup=ReplyKeyboardMarkup(
                    [["🔙 Главное меню"]],
                    one_time_keyboard=True
                )
            )
            context.user_data.clear()
            from commands.start_commands import exit_to_main_menu
            return await exit_to_main_menu(update, context)
        else:
            await update.message.reply_text(
                f"❌ Договор для ученика {telegram} не найден.\n"
                "Попробуйте еще раз или введите '🔙 Отмена':",
                reply_markup=ReplyKeyboardMarkup(
                    [["🔙 Отмена"]],
                    one_time_keyboard=True
                )
            )
            return CONTRACT_STUDENT_TG

    # Ищем студента - пробуем и с @, и без @
    student = session.query(Student).filter(
        (Student.telegram == telegram_clean) |
        (Student.telegram == telegram_with_at) |
        (Student.telegram == telegram)
    ).first()

    if not student:
        await update.message.reply_text(
            f"❌ Студент с Telegram {telegram} не найден.\n"
            "Попробуйте еще раз или введите '🔙 Отмена':"
        )
        return CONTRACT_STUDENT_TG

    # Сохраняем данные студента в context
    # Используем telegram из базы (может быть с @ или без)
    context.user_data['contract_student_id'] = student.id
    context.user_data['contract_student_telegram'] = student.telegram  # Используем telegram из базы
    context.user_data['contract_student_fio'] = student.fio

    # Показываем меню выбора типа договора
    keyboard = [
        ["Ручное", "Авто"],
        ["Фуллстек", "Экспресс Авто"],
        ["🔙 Отмена"]
    ]

    await update.message.reply_text(
        f"✅ Студент найден: {student.fio}\n\n"
        "Выберите тип договора:",
        reply_markup=ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
    )
    return CONTRACT_TYPE


async def handle_contract_type(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает выбор типа договора и переходит к выбору авансового платежа.
    """
    contract_type = update.message.text.strip()

    if contract_type == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    if contract_type not in CONTRACT_TEMPLATES:
        await update.message.reply_text(
            "❌ Неверный тип договора. Выберите из предложенных:",
            reply_markup=ReplyKeyboardMarkup(
                [["Ручное"], ["Авто"], ["Фуллстек"], ["Экспресс Авто"], ["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_TYPE

    context.user_data['contract_type'] = contract_type

    # Формируем клавиатуру с шаблонными суммами
    amounts = ADVANCE_AMOUNTS.get(contract_type, [])
    keyboard = []

    for amount_info in amounts:
        keyboard.append([f"{amount_info['amount']} руб. ({amount_info['text']})"])

    keyboard.append(["Ввести свою сумму"])
    keyboard.append(["🔙 Отмена"])

    await update.message.reply_text(
        f"📄 Тип договора: {contract_type}\n\n"
        "Выберите сумму авансового платежа:",
        reply_markup=ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
    )
    return CONTRACT_ADVANCE_AMOUNT


async def handle_advance_amount(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает выбор суммы авансового платежа.
    """
    text = update.message.text.strip()

    if text == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    if text == "Ввести свою сумму":
        await update.message.reply_text(
            "Введите сумму авансового платежа (только число, например: 50000):",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        context.user_data['waiting_custom_amount'] = True
        return CONTRACT_ADVANCE_AMOUNT

    # Если ожидаем ввод своей суммы
    if context.user_data.get('waiting_custom_amount'):
        try:
            amount = int(text)
            if amount <= 0:
                raise ValueError

            # Получаем пропись
            amount_text = number_to_words_rubles(amount).replace(" рублей", "").capitalize()

            context.user_data['advance_amount'] = amount
            context.user_data['advance_amount_text'] = amount_text
            context.user_data['waiting_custom_amount'] = False

            # Переходим к типу платежа
            await update.message.reply_text(
                f"✅ Авансовый платеж: {amount} руб. ({amount_text})\n\n"
                "Выберите тип платежа:",
                reply_markup=ReplyKeyboardMarkup(
                    [["Единоразовый платеж"], ["Ежемесячный платеж"], ["🔙 Отмена"]],
                    one_time_keyboard=True
                )
            )
            return CONTRACT_PAYMENT_TYPE

        except ValueError:
            await update.message.reply_text(
                "❌ Введите корректное число (например: 50000):",
                reply_markup=ReplyKeyboardMarkup(
                    [["🔙 Отмена"]],
                    one_time_keyboard=True
                )
            )
            return CONTRACT_ADVANCE_AMOUNT

    # Парсим сумму из текста
    try:
        # Извлекаем число из строки типа "46000 руб. (Сорок шесть тысяч)"
        amount_str = text.split()[0]
        amount = int(amount_str)

        # Получаем пропись
        contract_type = context.user_data.get('contract_type')
        amounts = ADVANCE_AMOUNTS.get(contract_type, [])
        amount_text = None

        for amount_info in amounts:
            if amount_info['amount'] == amount:
                amount_text = amount_info['text']
                break

        if not amount_text:
            amount_text = number_to_words_rubles(amount).replace(" рублей", "").capitalize()

        context.user_data['advance_amount'] = amount
        context.user_data['advance_amount_text'] = amount_text

    except (ValueError, IndexError):
        await update.message.reply_text(
            "❌ Неверный формат суммы. Введите число (например: 50000):",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_ADVANCE_AMOUNT

    # Спрашиваем тип платежа
    await update.message.reply_text(
        f"✅ Авансовый платеж: {amount} руб. ({context.user_data['advance_amount_text']})\n\n"
        "Выберите тип платежа:",
        reply_markup=ReplyKeyboardMarkup(
            [["Единоразовый платеж"], ["Ежемесячный платеж"], ["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_PAYMENT_TYPE


async def handle_payment_type(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает выбор типа платежа.
    """
    payment_type = update.message.text.strip()

    if payment_type == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    if payment_type not in ["Единоразовый платеж", "Ежемесячный платеж"]:
        await update.message.reply_text(
            "❌ Выберите тип платежа:",
            reply_markup=ReplyKeyboardMarkup(
                [["Единоразовый платеж"], ["Ежемесячный платеж"], ["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_PAYMENT_TYPE

    context.user_data['payment_type'] = payment_type

    if payment_type == "Ежемесячный платеж":
        await update.message.reply_text(
            "Введите количество месяцев:",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_MONTHS
    else:
        # Переходим к комиссии - определяем стандартную комиссию по типу договора
        contract_type = context.user_data.get('contract_type', '')
        if contract_type == "Фуллстек":
            default_commission = "2 месяца по 65%"
        elif contract_type == "Экспресс Авто":
            default_commission = "2 месяца по 25%"
        else:
            default_commission = "2 месяца по 55%"
        
        await update.message.reply_text(
            f"Выберите тип комиссии (п. 3.4):\n"
            f"Стандарт для {contract_type}: {default_commission}",
            reply_markup=ReplyKeyboardMarkup(
                [[default_commission], ["Уникальный"], ["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_COMMISSION_TYPE


async def handle_months(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод количества месяцев для ежемесячного платежа.
    """
    text = update.message.text.strip()

    if text == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    try:
        months = int(text)
        if months <= 0:
            raise ValueError
        context.user_data['payment_months'] = months
    except ValueError:
        await update.message.reply_text(
            "❌ Введите корректное количество месяцев (положительное число):",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_MONTHS

    # Переходим к комиссии - определяем стандартную комиссию по типу договора
    contract_type = context.user_data.get('contract_type', '')
    if contract_type == "Фуллстек":
        default_commission = "2 месяца по 65%"
    elif contract_type == "Экспресс Авто":
        default_commission = "2 месяца по 25%"
    else:
        default_commission = "2 месяца по 55%"
    
    await update.message.reply_text(
        f"✅ Количество месяцев: {months}\n\n"
        f"Выберите тип комиссии (п. 3.4):\n"
        f"Стандарт для {contract_type}: {default_commission}",
        reply_markup=ReplyKeyboardMarkup(
            [[default_commission], ["Уникальный"], ["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_COMMISSION_TYPE


async def handle_commission_type(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает выбор типа комиссии.
    """
    commission_type = update.message.text.strip()

    if commission_type == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    if commission_type == "2 месяца по 55%":
        context.user_data['commission_months'] = 2
        context.user_data['commission_percent'] = 55
    elif commission_type == "2 месяца по 65%":
        context.user_data['commission_months'] = 2
        context.user_data['commission_percent'] = 65
    elif commission_type == "2 месяца по 25%":
        context.user_data['commission_months'] = 2
        context.user_data['commission_percent'] = 25
    
    # Если выбрана стандартная комиссия, переходим к реквизитам
    if commission_type in ["2 месяца по 55%", "2 месяца по 65%", "2 месяца по 25%"]:
        # Переходим к реквизитам
        await update.message.reply_text(
            "Введите ФИО заказчика:",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_FIO
    elif commission_type == "Уникальный":
        await update.message.reply_text(
            "Введите комиссию в формате: количество_месяцев, процент\n"
            "Например: 2, 50 (2 месяца по 50%):",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_COMMISSION_CUSTOM
    else:
        await update.message.reply_text(
            "❌ Выберите тип комиссии:",
            reply_markup=ReplyKeyboardMarkup(
                [["2 месяца по 55%"], ["Уникальный"], ["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_COMMISSION_TYPE


async def handle_commission_custom(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод уникальной комиссии.
    """
    text = update.message.text.strip()

    if text == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    try:
        parts = text.split(",")
        if len(parts) != 2:
            raise ValueError

        months = int(parts[0].strip())
        percent = int(parts[1].strip())

        if months <= 0 or percent <= 0 or percent > 100:
            raise ValueError

        context.user_data['commission_months'] = months
        context.user_data['commission_percent'] = percent
    except (ValueError, IndexError):
        await update.message.reply_text(
            "❌ Неверный формат. Введите в формате: количество_месяцев, процент\n"
            "Например: 2, 55",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Отмена"]],
                one_time_keyboard=True
            )
        )
        return CONTRACT_COMMISSION_CUSTOM

    # Переходим к реквизитам
    await update.message.reply_text(
        f"✅ Комиссия: {months} месяцев по {percent}%\n\n"
        "Введите ФИО заказчика:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_FIO


async def handle_fio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод ФИО заказчика.
    """
    fio = update.message.text.strip()

    if fio == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_fio'] = fio

    await update.message.reply_text(
        f"✅ ФИО: {fio}\n\n"
        "Введите адрес:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_ADDRESS


async def handle_address(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод адреса.
    """
    address = update.message.text.strip()

    if address == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_address'] = address

    await update.message.reply_text(
        f"✅ Адрес: {address}\n\n"
        "Введите ИНН:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_INN


async def handle_inn(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод ИНН.
    """
    inn = update.message.text.strip()

    if inn == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_inn'] = inn

    await update.message.reply_text(
        f"✅ ИНН: {inn}\n\n"
        "Введите расчетный счет (р/с):",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_RS


async def handle_rs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод расчетного счета.
    """
    rs = update.message.text.strip()

    if rs == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_rs'] = rs

    await update.message.reply_text(
        f"✅ Расчетный счет: {rs}\n\n"
        "Введите корреспондентский счет (к/с):",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_KS


async def handle_ks(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод корреспондентского счета.
    """
    ks = update.message.text.strip()

    if ks == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_ks'] = ks

    await update.message.reply_text(
        f"✅ Корреспондентский счет: {ks}\n\n"
        "Введите название банка:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_BANK


async def handle_bank(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод названия банка.
    """
    bank = update.message.text.strip()

    if bank == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_bank'] = bank

    await update.message.reply_text(
        f"✅ Банк: {bank}\n\n"
        "Введите БИК:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_BIK


async def handle_bik(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод БИК.
    """
    bik = update.message.text.strip()

    if bik == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_bik'] = bik

    await update.message.reply_text(
        f"✅ БИК: {bik}\n\n"
        "Введите адрес электронной почты:",
        reply_markup=ReplyKeyboardMarkup(
            [["🔙 Отмена"]],
            one_time_keyboard=True
        )
    )
    return CONTRACT_EMAIL


async def handle_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Обрабатывает ввод email и формирует договор.
    """
    email = update.message.text.strip()

    if email == "🔙 Отмена":
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    context.user_data['contract_email'] = email

    # Формируем договор
    try:
        file_path = await generate_contract(context.user_data)

        # Отправляем файл пользователю
        with open(file_path, 'rb') as doc_file:
            await update.message.reply_document(
                document=doc_file,
                filename=os.path.basename(file_path),
                caption="✅ Договор успешно сформирован!"
            )

        await update.message.reply_text(
            "Договор сохранен и отправлен.\n"
            "Вы можете пересоздать договор или отправить его повторно.",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Главное меню"]],
                one_time_keyboard=True
            )
        )

        # Очищаем данные
        context.user_data.clear()

        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)

    except Exception as e:
        await update.message.reply_text(
            f"❌ Ошибка при формировании договора: {str(e)}\n"
            "Попробуйте еще раз.",
            reply_markup=ReplyKeyboardMarkup(
                [["🔙 Главное меню"]],
                one_time_keyboard=True
            )
        )
        from commands.start_commands import exit_to_main_menu
        return await exit_to_main_menu(update, context)


async def generate_contract(data: dict) -> str:
    """
    Генерирует договор на основе данных и возвращает путь к файлу.
    """
    contract_type = data['contract_type']
    template_path = CONTRACT_TEMPLATES[contract_type]
    
    # Получаем абсолютный путь относительно корня проекта
    base_dir = get_project_root()
    template_path = os.path.join(base_dir, template_path)
    
    # Логирование для отладки
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"Поиск шаблона: тип={contract_type}, base_dir={base_dir}, template_path={template_path}")
    logger.info(f"Проверка существования: {os.path.exists(template_path)}")
    
    # Проверяем, существует ли директория шаблонов
    templates_dir = os.path.join(base_dir, "doc", "шаблоны")
    logger.info(f"Директория шаблонов: {templates_dir}, существует: {os.path.exists(templates_dir)}")
    if os.path.exists(templates_dir):
        files = os.listdir(templates_dir)
        logger.info(f"Файлы в директории шаблонов: {files}")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон {template_path} не найден")

    # Открываем шаблон
    doc = Document(template_path)

    # Получаем данные для замены
    contract_date = get_contract_date_formatted()
    contract_number = get_contract_number()
    fio = data.get('contract_fio', '').strip()
    advance_amount = data.get('advance_amount', 0)
    advance_amount_text = data.get('advance_amount_text', '').strip()
    payment_type = data.get('payment_type', 'Единоразовый платеж')
    payment_months = data.get('payment_months', None)
    # Определяем стандартную комиссию по типу договора
    contract_type = data.get('contract_type', '')
    if contract_type == "Фуллстек":
        default_commission_percent = 65
    elif contract_type == "Экспресс Авто":
        default_commission_percent = 25
    else:
        default_commission_percent = 55
    
    commission_months = data.get('commission_months', 2)
    commission_percent = data.get('commission_percent', default_commission_percent)
    
    # Вычисляем суммарный процент (общий процент за все месяцы)
    total_commission_percent = commission_months * commission_percent

    # Реквизиты
    address = data.get('contract_address', '').strip()
    inn = data.get('contract_inn', '').strip()
    rs = data.get('contract_rs', '').strip()
    ks = data.get('contract_ks', '').strip()
    bank = data.get('contract_bank', '').strip()
    bik = data.get('contract_bik', '').strip()
    email = data.get('contract_email', '').strip()

    # Отладочная информация (можно убрать после тестирования)
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"Генерация договора: дата={contract_date}, номер={contract_number}, ФИО={fio}")
    logger.info(f"Реквизиты: адрес={address}, ИНН={inn}, р/с={rs}, к/с={ks}, банк={bank}, БИК={bik}, email={email}")

    # Формируем текст комиссии в нужном формате
    # Например: "50% в последний календарный день первого месяца со дня трудоустройства;\n50% в последний календарный день второго месяца со дня трудоустройства."
    commission_parts = []
    for month_num in range(1, commission_months + 1):
        # Определяем порядковое числительное
        month_names = {
            1: "первого",
            2: "второго",
            3: "третьего",
            4: "четвертого",
            5: "пятого",
            6: "шестого",
            7: "седьмого",
            8: "восьмого",
            9: "девятого",
            10: "десятого",
            11: "одиннадцатого",
            12: "двенадцатого"
        }
        month_name = month_names.get(month_num, f"{month_num}-го")
        commission_parts.append(
            f"{commission_percent}% в последний календарный день {month_name} месяца со дня трудоустройства;")

    commission_text = "\n".join(commission_parts)

    # Формируем текст платежа
    if payment_type == "Единоразовый платеж":
        payment_text = "единоразовый платеж"
        monthly_text = ""  # Убираем текст про ежемесячный
        # Универсальный текст для условий платежа
        payment_terms_text = "Заказчик обязан внести авансовый платеж Исполнителю за оказываемые услуги в течение 5 (пяти) рабочих дней со дня заключения Договора."
    else:
        payment_text = f"ежемесячный платеж на {payment_months} месяцев"
        monthly_text = f"на {payment_months} месяцев"
        # Универсальный текст для условий платежа (ежемесячный)
        # Формируем части динамически в зависимости от количества месяцев
        payment_parts = [
            "Заказчик вносит авансовый платеж Исполнителю за оказываемые услуги в течение {months} месяцев со дня заключения договора в следующем порядке:",
            "",
            "Часть первая: в течение 5 рабочих дней со дня заключения договора.",
            "",
            "Часть вторая: по истечении одного месяца со дня заключения договора."
        ]

        # Добавляем дополнительные части для месяцев >= 2
        if payment_months >= 2:
            # Определяем порядковые числительные для месяцев
            month_names = {
                2: "второго",
                3: "третьего",
                4: "четвертого",
                5: "пятого",
                6: "шестого",
                7: "седьмого",
                8: "восьмого",
                9: "девятого",
                10: "десятого",
                11: "одиннадцатого",
                12: "двенадцатого"
            }

            # Добавляем части для каждого месяца начиная со второго
            for month_num in range(2, payment_months + 1):
                month_name = month_names.get(month_num, f"{month_num}-го")
                part_number = month_num + 1  # Часть третья, четвертая и т.д.
                part_names = {
                    3: "третья",
                    4: "четвертая",
                    5: "пятая",
                    6: "шестая",
                    7: "седьмая",
                    8: "восьмая",
                    9: "девятая",
                    10: "десятая",
                    11: "одиннадцатая",
                    12: "двенадцатая"
                }
                part_name = part_names.get(part_number, f"{part_number}-я")
                payment_parts.append("")
                payment_parts.append(f"Часть {part_name}: по истечении {month_name} месяца со дня заключения договора.")

        # Заменяем {months} на количество месяцев
        payment_terms_text = "\n".join(payment_parts).replace("{months}", str(payment_months))

    # Словарь замен плейсхолдеров (соответствует плейсхолдерам в шаблоне)
    replacements = {
        # Основные поля
        "{{date}}": contract_date,
        "{{num_doc}}": contract_number,
        "{{fio}}": fio,
        "{{sum_num}}": str(advance_amount),
        "{{sum_word}}": advance_amount_text,

        # Платеж - универсальный плейсхолдер для условий платежа
        "{{payment_terms}}": payment_terms_text,
        "{{month}}": str(payment_months) if payment_months else "",

        # Комиссия
        "{{comission}}": commission_text,
        "{{procent}}": str(total_commission_percent),

        # Реквизиты
        "{{adress}}": address,
        "{{inn}}": inn,
        "{{r/s}}": rs,
        "{{k/s}}": ks,
        "{{bank}}": bank,
        "{{bic}}": bik,
        "{{email}}": email,

        # Для обратной совместимости (если в других шаблонах используются старые плейсхолдеры)
        "{{DATE}}": contract_date,
        "{{DATE}} г.": f"{contract_date} г.",
        "{{CONTRACT_NUMBER}}": contract_number,
        "{{FIO}}": fio,
        "{{ADVANCE_AMOUNT}}": str(advance_amount),
        "{{ADVANCE_AMOUNT_TEXT}}": advance_amount_text,
        "{{PAYMENT_TYPE}}": payment_text,
        "{{MONTHLY_MONTHS}}": monthly_text,
        "{{COMMISSION}}": commission_text,
        "{{PROCENT}}": str(total_commission_percent),
        "{{ADDRESS}}": address,
        "{{INN}}": inn,
        "{{RS}}": rs,
        "{{KS}}": ks,
        "{{BANK}}": bank,
        "{{BIK}}": bik,
        "{{EMAIL}}": email,
    }

    # Заменяем плейсхолдеры в документе
    # Важно: заменяем несколько раз, чтобы обработать все вхождения
    # Сначала основные поля (дата, номер, ФИО)
    priority_replacements = {
        "{{date}}": contract_date,
        "{{num_doc}}": contract_number,
        "{{fio}}": fio,
    }

    # Первый проход - основные поля
    for paragraph in doc.paragraphs:
        for old_text, new_text in priority_replacements.items():
            if new_text:
                replace_text_in_paragraph(paragraph, old_text, new_text)

    # Обрабатываем таблицы - основные поля
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in priority_replacements.items():
                        if new_text:
                            replace_text_in_paragraph(paragraph, old_text, new_text)

    # Второй проход - все остальные поля (делаем несколько проходов для надежности)
    # Отслеживаем обработанные параграфы для payment_terms, чтобы избежать дублирования
    processed_paragraphs_payment_terms = set()
    
    for _ in range(3):  # 3 прохода для обработки всех вхождений
        for paragraph in doc.paragraphs:
            paragraph_id = id(paragraph)
            for old_text, new_text in replacements.items():
                # Пропускаем пустые замены и уже обработанные
                if new_text == "" or old_text in priority_replacements:
                    continue
                # Для payment_terms проверяем, не был ли уже обработан этот параграф
                if old_text == "{{payment_terms}}":
                    if paragraph_id in processed_paragraphs_payment_terms:
                        continue
                    if replace_text_in_paragraph(paragraph, old_text, new_text):
                        processed_paragraphs_payment_terms.add(paragraph_id)
                else:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

        # Обрабатываем таблицы - все остальные поля
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_id = id(paragraph)
                        for old_text, new_text in replacements.items():
                            # Пропускаем пустые замены и уже обработанные
                            if new_text == "" or old_text in priority_replacements:
                                continue
                            # Для payment_terms проверяем, не был ли уже обработан этот параграф
                            if old_text == "{{payment_terms}}":
                                if paragraph_id in processed_paragraphs_payment_terms:
                                    continue
                                if replace_text_in_paragraph(paragraph, old_text, new_text):
                                    processed_paragraphs_payment_terms.add(paragraph_id)
                            else:
                                replace_text_in_paragraph(paragraph, old_text, new_text)

    # Обрабатываем заголовки и футеры
    for section in doc.sections:
        # Заголовки
        for header in section.header.paragraphs:
            header_id = id(header)
            for old_text, new_text in replacements.items():
                if new_text:
                    # Для payment_terms проверяем, не был ли уже обработан этот параграф
                    if old_text == "{{payment_terms}}":
                        if header_id not in processed_paragraphs_payment_terms:
                            if replace_text_in_paragraph(header, old_text, new_text):
                                processed_paragraphs_payment_terms.add(header_id)
                    else:
                        replace_text_in_paragraph(header, old_text, new_text)

        # Футеры
        for footer in section.footer.paragraphs:
            footer_id = id(footer)
            for old_text, new_text in replacements.items():
                if new_text:
                    # Для payment_terms проверяем, не был ли уже обработан этот параграф
                    if old_text == "{{payment_terms}}":
                        if footer_id not in processed_paragraphs_payment_terms:
                            if replace_text_in_paragraph(footer, old_text, new_text):
                                processed_paragraphs_payment_terms.add(footer_id)
                    else:
                        replace_text_in_paragraph(footer, old_text, new_text)

    # Обработка универсального плейсхолдера {{payment_terms}} уже выполнена через словарь replacements выше
    # Также обрабатываем старые плейсхолдеры для обратной совместимости (если они остались в шаблоне)
    old_placeholders_to_remove = ["{{one-time payment}}", "{{Monthly_payment}}"]

    for paragraph in doc.paragraphs:
        full_text = ''.join([run.text for run in paragraph.runs])
        # Удаляем старые плейсхолдеры, если они остались
        for old_ph in old_placeholders_to_remove:
            if old_ph in full_text:
                new_text = full_text.replace(old_ph, "")
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new_text

    # То же для таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join([run.text for run in paragraph.runs])
                    for old_ph in old_placeholders_to_remove:
                        if old_ph in full_text:
                            new_text = full_text.replace(old_ph, "")
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = new_text

    # Сохраняем файл
    student_telegram = data.get('contract_student_telegram', 'unknown')
    filename = f"{student_telegram}.docx"
    
    # Получаем абсолютный путь относительно корня проекта
    base_dir = get_project_root()
    doc_dir = os.path.join(base_dir, "doc")
    file_path = os.path.join(doc_dir, filename)

    # Создаем папку если не существует
    os.makedirs(doc_dir, exist_ok=True)

    doc.save(file_path)
    return file_path


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Заменяет текст в параграфе, сохраняя форматирование.
    Обрабатывает случаи, когда плейсхолдер может быть разбит на несколько runs.
    """
    if not old_text or not new_text:
        return False

    # Собираем весь текст из всех runs
    runs = paragraph.runs
    if not runs:
        return False

    full_text = ''.join([run.text for run in runs])

    # Проверяем, есть ли плейсхолдер в тексте
    if old_text not in full_text:
        return False

    # Заменяем все вхождения
    new_full_text = full_text.replace(old_text, new_text)

    # Если текст не изменился, выходим
    if new_full_text == full_text:
        return False

    # Очищаем все runs
    for run in runs:
        run.text = ""

    # Добавляем новый текст в первый run (сохраняем форматирование первого run)
    if runs:
        runs[0].text = new_full_text
    else:
        paragraph.add_run(new_full_text)

    return True

